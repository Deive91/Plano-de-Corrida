import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_plano_doc():
    doc = docx.Document()

    # Configurar margens
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Estilos de Cores
    COR_TITULO = RGBColor(16, 185, 129)    # Emerald
    COR_SUBTITULO = RGBColor(245, 158, 11) # Amber
    COR_TEXTO = RGBColor(51, 65, 85)       # Slate 700

    def add_title(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = COR_TITULO
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COR_TITULO
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COR_SUBTITULO
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = 'Segoe UI'
            run_b.font.size = Pt(11)
            run_b.font.bold = True
            run_b.font.color.rgb = RGBColor(30, 41, 59)
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.font.color.rgb = COR_TEXTO
        return p

    # TÍTULO
    add_title("DOCUMENTAÇÃO TÉCNICA E CIENTÍFICA\nPLANO DE CORRIDA APP")
    add_p("Data: Julho de 2026 | Versão 2.6 (Esvaziamento do Catálogo de Treinos) | Arquitetura: React 18 + Vite + JSON API")

    add_h1("1. Esvaziamento e Limpeza do Catálogo (v2.6)")
    add_p("Na versão 2.6, conforme solicitado, todos os treinos do catálogo foram apagados (workouts: []) em workouts_api_catalog.json. O sistema mantém sua arquitetura limpa e pronta para receber uma nova alimentação de dados JSON sem gerar erros na interface ou na compilação.")

    add_h1("2. Estrutura Base de Fallback do Motor de Treinamento")
    add_p("Enquanto o banco de dados JSON estiver vazio, o motor de treinamento (trainingEngine.js) utiliza as descrições limpas e fisiológicas de fallback base, garantindo que o cronograma semanal continue funcional para o corredor com orientações contínuas para cada tipo de sessão (LONG, EASY, INTERVAL, TEMPO, HILL, REST e STRENGTH).")

    add_h1("3. Regra Universal: Fortalecimento Opcional nas Segundas-feiras")
    add_p("A Segunda-feira de todos os planos permanece configurada com o Treino de Fortalecimento Funcional e Mobilidade com marcação de ⭐ Opcional (`isOptional: true`). Dessa forma, o atleta obtém suporte biomecânico contra lesões caso treine, ou descanso completo sem perda de porcentagem no progresso semanal.")

    out_path = os.path.join(os.path.dirname(__file__), "Documentacao_PlanoDeCorrida.docx")
    doc.save(out_path)
    print(f"Documento DOCX salvo em: {out_path}")

if __name__ == "__main__":
    create_plano_doc()
